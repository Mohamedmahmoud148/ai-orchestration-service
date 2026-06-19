"""
app/modules/file_extraction.py  —  v3.0

Smart File Extraction Module.

Supports:
  - PDF  (text-based via pdfminer → fallback pypdf → fallback vision model)
  - DOCX (python-docx)
  - XLSX / XLS (openpyxl / xlrd) — full table extraction with headers
  - Images (PNG/JPG) — vision model via OpenRouter
  - TXT / CSV

File source priority:
  1. file_bytes  in context
  2. file_url    in context (direct download)
  3. file_reference via backend API
  4. subjectOfferingId → fetch first material from offering
"""
from __future__ import annotations

import base64
import io
from typing import Any, Dict, List, Optional

import httpx

from app.agents.schemas import AgentInput, AgentOutput
from app.core.logging import logger

# ── Vision model for scanned PDFs / images ────────────────────────────────────
_VISION_MODEL = "google/gemini-flash-1.5"   # via OpenRouter — fast + cheap + great OCR
_VISION_MAX_PAGES = 5                        # max PDF pages to send as images

# ── Excel column limit — avoid token explosion ────────────────────────────────
_MAX_EXCEL_ROWS = 200
_MAX_EXCEL_COLS = 30


# ─────────────────────────────────────────────────────────────────────────────
# TEXT EXTRACTORS
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf_text(data: bytes) -> str:
    """Try pdfminer first, then pypdf. Returns empty string if both fail."""
    # pdfminer (best quality)
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(io.BytesIO(data)) or ""
        if text.strip():
            return text
    except Exception as e:
        logger.debug("pdfminer failed: %s", e)

    # pypdf fallback
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(p for p in pages if p.strip())
        return text
    except Exception as e:
        logger.debug("pypdf failed: %s", e)

    return ""


def _extract_docx(data: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        parts: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning("DOCX extraction failed: %s", e)
        return ""


def _extract_excel(data: bytes, filename: str = "") -> str:
    """
    Extract Excel file (xlsx or xls) into a readable text table.
    Returns all sheets with headers and data rows.
    """
    is_xls = filename.lower().endswith(".xls") and not filename.lower().endswith(".xlsx")

    if is_xls:
        return _extract_xls(data)

    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        output: List[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(max_row=_MAX_EXCEL_ROWS, max_col=_MAX_EXCEL_COLS, values_only=True))

            if not rows:
                continue

            # First row as headers
            headers = [str(h) if h is not None else "" for h in rows[0]]
            output.append(f"📊 Sheet: {sheet_name}")
            output.append(" | ".join(headers))
            output.append("-" * min(80, len(" | ".join(headers))))

            for row in rows[1:]:
                cells = [str(c) if c is not None else "" for c in row]
                # Skip completely empty rows
                if any(c.strip() for c in cells):
                    output.append(" | ".join(cells))

            if len(list(ws.iter_rows())) > _MAX_EXCEL_ROWS:
                output.append(f"... (showing first {_MAX_EXCEL_ROWS} rows)")

            output.append("")  # blank line between sheets

        wb.close()
        return "\n".join(output)

    except Exception as e:
        logger.warning("openpyxl failed: %s", e)
        return _extract_xls(data)  # fallback to xlrd


def _extract_xls(data: bytes) -> str:
    try:
        import xlrd
        wb = xlrd.open_workbook(file_contents=data)
        output: List[str] = []

        for sheet in wb.sheets():
            output.append(f"📊 Sheet: {sheet.name}")
            for rx in range(min(sheet.nrows, _MAX_EXCEL_ROWS)):
                row_vals = [str(sheet.cell_value(rx, cx)) for cx in range(min(sheet.ncols, _MAX_EXCEL_COLS))]
                output.append(" | ".join(row_vals))
            output.append("")

        return "\n".join(output)
    except Exception as e:
        logger.warning("xlrd failed: %s", e)
        return ""


def _extract_csv(data: bytes) -> str:
    try:
        import csv
        text = data.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(text))
        rows = [" | ".join(row) for row in reader]
        return "\n".join(rows[:_MAX_EXCEL_ROWS])
    except Exception as e:
        logger.warning("CSV extraction failed: %s", e)
        return data.decode("utf-8", errors="ignore")[:5000]


def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Dispatch extraction based on file extension."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf_text(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".xlsx") or name.endswith(".xls"):
        return _extract_excel(data, filename)
    if name.endswith(".csv"):
        return _extract_csv(data)
    if name.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
        return ""  # images need vision — handled by caller
    return _extract_txt(data)


# ─────────────────────────────────────────────────────────────────────────────
# VISION (for scanned PDFs + images)
# ─────────────────────────────────────────────────────────────────────────────

def _pdf_pages_to_images(data: bytes, max_pages: int = _VISION_MAX_PAGES) -> List[bytes]:
    """Convert PDF pages to PNG images using pypdf + Pillow."""
    images: List[bytes] = []
    try:
        import pypdf
        from PIL import Image as PILImage

        reader = pypdf.PdfReader(io.BytesIO(data))
        for i, page in enumerate(reader.pages[:max_pages]):
            # Try to extract images from the page
            for img_obj in page.images:
                buf = io.BytesIO()
                try:
                    pil_img = PILImage.open(io.BytesIO(img_obj.data))
                    pil_img.save(buf, format="PNG")
                    images.append(buf.getvalue())
                    break  # one image per page is enough
                except Exception:
                    pass
    except Exception as e:
        logger.debug("PDF-to-image failed: %s", e)
    return images


async def _vision_extract(
    data: bytes,
    filename: str,
    user_question: str,
    model_router: Any,
    auth_header: Optional[str] = None,
) -> str:
    """
    Send file to vision model (via OpenRouter) for OCR + understanding.
    Works for: scanned PDFs, images, any file where text extraction returns empty.
    """
    try:
        name = (filename or "").lower()
        image_bytes_list: List[bytes] = []

        # For images, use directly
        if name.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
            image_bytes_list = [data]
        # For PDFs, try to convert pages to images
        elif name.endswith(".pdf"):
            image_bytes_list = _pdf_pages_to_images(data)

        if not image_bytes_list:
            return ""

        # Build vision messages for OpenRouter
        content_parts: List[dict] = [
            {
                "type": "text",
                "text": (
                    f"Please read and extract ALL content from this document.\n"
                    f"File: {filename}\n"
                    f"User question: {user_question}\n\n"
                    "Extract all text, tables, numbers, and data. "
                    "If it's a table, format it clearly with columns. "
                    "Respond in the same language as the user's question."
                )
            }
        ]

        for img_bytes in image_bytes_list[:_VISION_MAX_PAGES]:
            b64 = base64.b64encode(img_bytes).decode()
            content_parts.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{b64}"}
            })

        messages = [
            {"role": "system", "content": "You are an expert document reader with OCR capabilities. Extract content accurately."},
            {"role": "user", "content": content_parts}
        ]

        result = await model_router.generate_with_messages(
            messages=messages,
            model_id=_VISION_MODEL,
        )
        return result or ""

    except Exception as e:
        logger.warning("Vision extraction failed: %s", e)
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# FILE FETCHER
# ─────────────────────────────────────────────────────────────────────────────

async def _fetch_from_url(url: str, auth_header: Optional[str]) -> Optional[bytes]:
    """Download file bytes from a URL."""
    try:
        headers = {}
        if auth_header:
            headers["Authorization"] = auth_header
        async with httpx.AsyncClient() as client:
            resp = await client.get(url, headers=headers, timeout=30.0, follow_redirects=True)
            resp.raise_for_status()
        return resp.content
    except Exception as e:
        logger.warning("URL fetch failed (%s): %s", url[:80], e)
        return None


# ─────────────────────────────────────────────────────────────────────────────
# MAIN MODULE
# ─────────────────────────────────────────────────────────────────────────────

class FileExtractionModule:
    """
    Smart file extraction + understanding.

    Pipeline:
      1. Get bytes (from context / URL / backend API / offering materials)
      2. Extract text (PDF → pdfminer/pypdf, Excel → openpyxl, DOCX, CSV, TXT)
      3. If text extraction fails (scanned PDF / image) → Vision model OCR
      4. LLM understands + answers the user's question based on extracted content
    """

    def __init__(self, model_router, backend_client):
        self.model_router   = model_router
        self.backend_client = backend_client

    async def run(self, agent_input: AgentInput, plan=None) -> AgentOutput:
        logger.info("FileExtractionModule: starting.")

        ctx   = agent_input.context or {}
        msg   = agent_input.message or ""

        file_bytes: Optional[bytes]  = ctx.get("file_bytes")
        file_name: str               = ctx.get("file_name", "document.pdf")
        file_reference: Optional[str] = ctx.get("file_reference")

        # Check academic_context too (restored from memory by Agent)
        academic_ctx = ctx.get("academic_context", {}) or {}
        file_url: Optional[str] = (
            ctx.get("file_url") or ctx.get("fileUrl")
            or ctx.get("signedUrl") or ctx.get("url")
            or academic_ctx.get("file_url") or academic_ctx.get("fileUrl")
        )
        if not file_name or file_name == "document.pdf":
            file_name = (
                ctx.get("file_name") or academic_ctx.get("file_name", "document.pdf")
            )

        # ── 1. Get file bytes ─────────────────────────────────────────────────

        # Source A: direct bytes already in context
        if not file_bytes and file_reference:
            result = await self.backend_client.fetch(
                route=f"/api/Files/{file_reference}",
                auth_header=agent_input.auth_header,
            )
            if "error" not in result:
                file_bytes = result.get("_raw_bytes")
                file_name  = result.get("fileName", file_name)

        # Source B: direct URL
        if not file_bytes and file_url and file_url.startswith("http"):
            file_bytes = await _fetch_from_url(file_url, agent_input.auth_header)
            url_path = file_url.split("?")[0]
            if "/" in url_path:
                file_name = url_path.split("/")[-1] or file_name

        # Source C: fetch first material from subjectOfferingId
        if not file_bytes:
            offering_id = (
                ctx.get("subjectOfferingId")
                or ctx.get("academic_context", {}).get("subjectOfferingId")
            )
            if offering_id:
                result = await self.backend_client.fetch(
                    route=f"/api/Materials/by-offering/{offering_id}",
                    auth_header=agent_input.auth_header,
                )
                items = result.get("items") or (result if isinstance(result, list) else [])
                if items:
                    first = items[0] if isinstance(items, list) else {}
                    mat_url = (
                        first.get("fileUrl") or first.get("signedUrl")
                        or first.get("url") or first.get("filePath") or ""
                    )
                    file_name = first.get("fileName") or first.get("title") or file_name
                    if mat_url and mat_url.startswith("http"):
                        file_bytes = await _fetch_from_url(mat_url, agent_input.auth_header)

        if not file_bytes:
            return AgentOutput(
                status="failed",
                response=(
                    "مش لاقي أي ملف أقرأه.\n"
                    "تأكد إن الدكتور رفع المواد الدراسية، ثم حاول مجدداً.\n\n"
                    "No file found. Please ensure course materials are uploaded first."
                ),
            )

        logger.info("FileExtractionModule: got %d bytes — '%s'", len(file_bytes), file_name)

        # ── 2. Extract text ───────────────────────────────────────────────────

        raw_text = extract_text_from_bytes(file_bytes, file_name)
        used_vision = False

        # ── 3. Vision fallback if text extraction empty ───────────────────────
        name_lower = file_name.lower()
        needs_vision = (
            not raw_text.strip()
            and (
                name_lower.endswith((".pdf", ".png", ".jpg", ".jpeg", ".gif", ".webp"))
            )
        )

        if needs_vision:
            logger.info("FileExtractionModule: text empty — trying vision model '%s'", _VISION_MODEL)
            raw_text = await _vision_extract(
                file_bytes, file_name, msg, self.model_router, agent_input.auth_header
            )
            used_vision = bool(raw_text.strip())

            if not raw_text.strip():
                return AgentOutput(
                    status="failed",
                    response=(
                        f"مش قادر أستخرج محتوى الملف '{file_name}'.\n"
                        "يمكن الملف تالف أو مش قابل للقراءة. "
                        "جرّب رفعه تاني بصيغة مختلفة."
                    ),
                )

        logger.info(
            "FileExtractionModule: extracted %d chars — vision=%s",
            len(raw_text), used_vision
        )

        # ── 4. LLM understands content + answers user's question ──────────────
        role = ctx.get("role", "student")

        system_prompt = (
            "أنت خبير في قراءة وتحليل الوثائق الجامعية.\n"
            "لديك المحتوى الكامل للملف أدناه.\n"
            "أجب على سؤال المستخدم بناءً على هذا المحتوى فقط.\n"
            "إذا كان السؤال عاماً (مثل 'اقرأ الملف') → لخّص المحتوى بشكل منظم.\n"
            "استخدم نفس لغة المستخدم (عربي → عربي، إنجليزي → إنجليزي).\n"
            "إذا كانت بيانات جدولية (Excel) → اعرضها بشكل واضح مع العناوين.\n"
            "لا تخترع معلومات غير موجودة في الملف."
        )

        file_type_label = _get_file_type_label(file_name)

        # ── Deep PDF mode for large documents ────────────────────────────────
        is_pdf = file_name.lower().endswith(".pdf")
        if is_pdf and len(raw_text) > 8_000:
            from app.modules.deep_pdf_study import extract_pdf_pages, summarize_large_pdf
            pages = extract_pdf_pages(file_bytes)
            if pages:
                logger.info("FileExtractionModule: deep PDF mode — %d pages", len(pages))
                explanation = await summarize_large_pdf(
                    pages=pages,
                    user_question=msg or "اشرح محتوى الملف",
                    model_router=self.model_router,
                    model_id="openai/gpt-4o-mini",
                )
                return AgentOutput(
                    status="success",
                    response=explanation or raw_text[:3000],
                    data={
                        "module":     "FileExtractionModule",
                        "file_name":  file_name,
                        "file_type":  file_type_label,
                        "page_count": len(pages),
                        "char_count": len(raw_text),
                        "mode":       "deep_pdf",
                    },
                )

        # ── Short document or non-PDF: single LLM call ────────────────────────
        # Cap at 30K chars — enough for ~15 pages without overflowing context
        content_preview = raw_text[:30_000]

        user_prompt = (
            f"نوع الملف: {file_type_label}\n"
            f"اسم الملف: {file_name}\n"
            f"{'(تم استخراجه بالرؤية الحاسوبية)' if used_vision else ''}\n\n"
            f"محتوى الملف:\n{content_preview}\n\n"
            f"سؤال/طلب المستخدم: {msg or 'لخص محتوى الملف'}"
        )

        answer = await self.model_router.generate(
            prompt=user_prompt,
            system_instruction=system_prompt,
            model_id="openai/gpt-4o-mini",
        )

        return AgentOutput(
            status="success",
            response=answer or raw_text[:2000],
            data={
                "module":        "FileExtractionModule",
                "file_name":     file_name,
                "file_type":     file_type_label,
                "char_count":    len(raw_text),
                "used_vision":   used_vision,
                "vision_model":  _VISION_MODEL if used_vision else None,
                "raw_preview":   raw_text[:300],
            },
        )


def _get_file_type_label(filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):    return "PDF"
    if name.endswith(".docx"):   return "Word Document"
    if name.endswith(".xlsx"):   return "Excel Spreadsheet"
    if name.endswith(".xls"):    return "Excel Spreadsheet (Legacy)"
    if name.endswith(".csv"):    return "CSV Table"
    if name.endswith(".txt"):    return "Text File"
    if name.endswith((".png", ".jpg", ".jpeg")): return "Image"
    return "Document"
